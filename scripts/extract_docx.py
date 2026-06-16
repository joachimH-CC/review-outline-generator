"""
Extract text and embedded images from Word (.docx) files.
Usage: python extract_docx.py <docx_dir> <output_text_dir> <output_image_dir>
"""
import docx
import os
import sys


def extract_docx(docx_path, text_dir, image_dir):
    """Extract text and images from a single docx file."""
    doc = docx.Document(docx_path)
    base_name = os.path.splitext(os.path.basename(docx_path))[0]

    # Extract text
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())

    # Also extract text from tables (preserve table structure)
    for table_idx, table in enumerate(doc.tables):
        header = []
        for row_idx, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                if row_idx == 0:
                    # Header row
                    header = row_text
                    texts.append(f'[Table {table_idx+1}] ' + ' | '.join(row_text))
                    texts.append('|' + '|'.join(['---'] * len(row_text)) + '|')
                else:
                    texts.append('| ' + ' | '.join(row_text) + ' |')

    text_out = os.path.join(text_dir, f"{base_name}.txt")
    with open(text_out, 'w', encoding='utf-8') as f:
        f.write('\n'.join(texts))

    # Extract images
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image = rel.target_part
            ext = os.path.splitext(image.partname)[1]
            if ext.lower() in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'):
                img_name = f"{base_name}_{image_count}{ext}"
                img_path = os.path.join(image_dir, img_name)
                with open(img_path, 'wb') as f:
                    f.write(image.blob)
                image_count += 1

    return len(texts), image_count


def main():
    if len(sys.argv) < 4:
        print("Usage: python extract_docx.py <docx_dir> <text_dir> <image_dir>")
        sys.exit(1)

    docx_dir = sys.argv[1]
    text_dir = sys.argv[2]
    image_dir = sys.argv[3]

    os.makedirs(text_dir, exist_ok=True)
    os.makedirs(image_dir, exist_ok=True)

    for f in sorted(os.listdir(docx_dir)):
        if f.endswith('.docx') and not f.startswith('~'):
            path = os.path.join(docx_dir, f)
            text_len, img_count = extract_docx(path, text_dir, image_dir)
            print(f"[docx] {f}: {text_len} text blocks, {img_count} images")


if __name__ == '__main__':
    main()
